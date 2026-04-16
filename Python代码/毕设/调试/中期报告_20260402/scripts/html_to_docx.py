#!/usr/bin/env python3
import sys
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
CODE_ROOT = SCRIPT_DIR.parent
VENDOR = CODE_ROOT / ".vendor"
sys.path.insert(0, str(VENDOR))

from bs4 import BeautifulSoup, NavigableString  # type: ignore
from docx import Document  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.shared import Pt  # type: ignore


REPORT_ROOT = Path("/Users/jia/Desktop/学习 /毕业设计/中期报告_20260402/report")
HTML_PATH = REPORT_ROOT / "中期报告_何嘉_20260402.html"
DOCX_PATH = REPORT_ROOT / "中期报告_何嘉_20260402.docx"


def set_fonts(run, size=12, bold=False):
    run.font.name = "Songti SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def append_text(paragraph, text, size=12, bold=False):
    run = paragraph.add_run(text)
    set_fonts(run, size=size, bold=bold)


def add_paragraph(document, text, size=12, bold=False, align=None, indent=True):
    p = document.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    append_text(p, text, size=size, bold=bold)
    return p


def clean_text(node):
    if isinstance(node, NavigableString):
        return str(node).strip()
    return node.get_text(" ", strip=True)


def html_table_to_docx(document, table_tag):
    rows = table_tag.find_all("tr")
    if not rows:
        return

    parsed_rows = []
    max_cols = 0
    for tr in rows:
        cells = tr.find_all(["th", "td"])
        parsed = []
        for cell in cells:
            colspan = int(cell.get("colspan", 1))
            parsed.append((clean_text(cell), colspan))
            max_cols += 0
        parsed_rows.append(parsed)
        max_cols = max(max_cols, sum(colspan for _, colspan in parsed))

    table = document.add_table(rows=len(parsed_rows), cols=max_cols)
    table.style = "Table Grid"

    for i, row in enumerate(parsed_rows):
        col_idx = 0
        for text, colspan in row:
            cell = table.cell(i, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            set_fonts(run, size=10.5, bold=(i == 0))
            if colspan > 1:
                end_cell = table.cell(i, col_idx + colspan - 1)
                cell.merge(end_cell)
            col_idx += colspan


def main():
    soup = BeautifulSoup(HTML_PATH.read_text(encoding="utf-8"), "html.parser")
    body = soup.body
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Songti SC"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)

    for node in body.children:
        if isinstance(node, NavigableString):
            continue
        name = getattr(node, "name", None)
        if not name:
            continue
        text = clean_text(node)
        if not text and name != "table":
            continue

        if name == "h1":
            add_paragraph(document, text, size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
        elif name == "h2":
            add_paragraph(document, text, size=13, bold=True, indent=False)
        elif name == "h3":
            add_paragraph(document, text, size=12, bold=True, indent=False)
        elif name == "p":
            add_paragraph(document, text, size=12, bold=False, indent=("note" not in node.get("class", [])))
        elif name == "table":
            html_table_to_docx(document, node)
        elif name == "div":
            document.add_paragraph()

    document.save(DOCX_PATH)


if __name__ == "__main__":
    main()
