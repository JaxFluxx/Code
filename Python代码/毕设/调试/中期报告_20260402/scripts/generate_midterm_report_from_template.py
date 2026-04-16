#!/usr/bin/env python3
import csv
import sys
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
CODE_ROOT = SCRIPT_DIR.parent
VENDOR = CODE_ROOT / ".vendor"
sys.path.insert(0, str(VENDOR))

from docx import Document  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.shared import Pt, Cm  # type: ignore


PROJECT_ROOT = Path("/Users/jia/Desktop/学习 /毕业设计")
REPORT_PATH = PROJECT_ROOT / "中期报告_20260402" / "report" / "中期报告_何嘉_20260402.docx"
TEMPLATE_PATH = PROJECT_ROOT / "开题答辩" / "IMG_F60FB994A4A6-1.docx"
SAMPLE_SUMMARY = PROJECT_ROOT / "实验" / "CODEX_clean_baseline" / "sample_summary.tsv"
PAIR_MANIFEST = PROJECT_ROOT / "实验" / "CODEX_clean_baseline" / "pair_manifest.tsv"
STRATIFIED_STATS = PROJECT_ROOT / "实验" / "CODEX_pair_results_r20_r30_stratified" / "stratified_stats.tsv"
COHORT_ANALYSIS = PROJECT_ROOT / "实验" / "CODEX_pair_results_r20_r30_stratified" / "cohort_analysis.tsv"


def read_tsv(path):
    with path.open("r", encoding="utf-8") as f:
        return list(csv.DictReader(f, delimiter="\t"))


def set_run_font(run, size=12, bold=False):
    run.font.name = "Songti SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_text(cell, text, size=10.5, bold=False, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def append_paragraph(cell, text, size=12, bold=False, first_indent=True, center=False):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def clear_cell(cell):
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)
    tc.append(OxmlElement("w:p"))


def ensure_full_width_body_cell(cell):
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        tc_pr = OxmlElement("w:tcPr")
        cell._tc.insert(0, tc_pr)

    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), "9185")
    tc_w.set(qn("w:type"), "dxa")

    grid_span = tc_pr.find(qn("w:gridSpan"))
    if grid_span is None:
        grid_span = OxmlElement("w:gridSpan")
        tc_pr.append(grid_span)
    grid_span.set(qn("w:val"), "6")


def fill_meta(table):
    set_cell_text(table.cell(0, 1), "数据科学与\n大数据技术", center=True)
    set_cell_text(table.cell(0, 3), "10042241", center=True)
    set_cell_text(table.cell(0, 5), "1004224119", center=True)

    set_cell_text(table.cell(1, 1), "何嘉", center=True)
    set_cell_text(table.cell(1, 3), "何勇强", center=True)
    set_cell_text(table.cell(1, 5), "高级工程师", center=True)

    set_cell_text(table.cell(2, 0), "题目类型", center=True)
    set_cell_text(table.cell(2, 1), "毕业设计", center=True)
    set_cell_text(table.cell(2, 3), "题目来源", center=True)
    set_cell_text(table.cell(2, 4), "科学研究", center=True)

    title = "基于CODEX空间蛋白组学的胃癌卵巢转移肿瘤微环境空间结构计算分析"
    set_cell_text(table.cell(3, 0), "设计（论文）题目", center=True)
    set_cell_text(table.cell(3, 1), title, center=True)


def get_metric(rows, group, metric):
    for row in rows:
        if row["group"] == group and row["metric"] == metric:
            return row
    return None


def format_num(value, digits=2):
    if value is None:
        return ""
    try:
        num = float(value)
    except (TypeError, ValueError):
        return str(value)
    if num.is_integer():
        return str(int(num))
    text = f"{num:.{digits}f}"
    return text.rstrip("0").rstrip(".")


def add_simple_table(cell, headers, rows):
    table = cell.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, size=10.5, bold=True, center=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), size=10.5, center=True)
    format_three_line_table(table)
    return table


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge)
        tag = qn(f"w:{edge}")
        element = tc_borders.find(tag)
        if edge_data:
            if element is None:
                element = OxmlElement(f"w:{edge}")
                tc_borders.append(element)
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))
        elif element is not None:
            tc_borders.remove(element)


def format_three_line_table(table):
    n_rows = len(table.rows)
    n_cols = len(table.columns)
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            borders = {
                "left": {"val": "nil"},
                "right": {"val": "nil"},
                "top": {"val": "nil"},
                "bottom": {"val": "nil"},
            }
            if r_idx == 0:
                borders["top"] = {"val": "single", "sz": "12", "color": "000000", "space": "0"}
                borders["bottom"] = {"val": "single", "sz": "8", "color": "000000", "space": "0"}
            if r_idx == n_rows - 1:
                borders["bottom"] = {"val": "single", "sz": "12", "color": "000000", "space": "0"}
            set_cell_border(cell, **borders)
    if n_cols == 3:
        widths = [Cm(3.0), Cm(8.0), Cm(4.0)]
    elif n_cols == 6:
        widths = [Cm(2.2), Cm(4.0), Cm(2.0), Cm(2.2), Cm(2.2), Cm(4.4)]
    elif n_cols == 7:
        widths = [Cm(1.5), Cm(2.8), Cm(1.3), Cm(1.7), Cm(1.8), Cm(1.6), Cm(4.0)]
    else:
        widths = [Cm(16.0 / max(n_cols, 1))] * n_cols
    for row in table.rows:
        for idx, width in enumerate(widths[: len(row.cells)]):
            row.cells[idx].width = width


def build_body(cell):
    clear_cell(cell)
    ensure_full_width_body_cell(cell)

    sample_rows = read_tsv(SAMPLE_SUMMARY)
    pair_rows = read_tsv(PAIR_MANIFEST)
    stratified_rows = read_tsv(STRATIFIED_STATS)
    cohort_rows = read_tsv(COHORT_ANALYSIS)

    append_paragraph(cell, "1.进度情况", bold=True, first_indent=False)
    append_paragraph(
        cell,
        "目前工作总体与开题阶段的时间安排基本一致，已完成数据整理、样本配对、肿瘤细胞与巨噬细胞baseline构建、空间指标计算以及初步分层统计分析。根据现有实验结果，中期阶段的重点已从“能否完成分析”转为“如何稳定主分析口径并收敛结论”。",
    )
    append_paragraph(
        cell,
        "与开题报告第1.2节提出的研究目标相比，当前已实际落地的是“建立具有明确生物学含义的空间指标体系并完成原发灶与卵巢转移灶的配对比较”；而开题中作为拓展方向提出的无监督聚类分析，目前仍保留在后续补充阶段，尚未作为中期主体内容展开。",
    )
    append_paragraph(cell, "当前进度时间线如下：", first_indent=False)
    add_simple_table(
        cell,
        ["时间", "阶段任务", "当前完成情况"],
        [
            ["2026年1月", "开题立项、数据说明梳理、样本编码核对", "已完成"],
            ["2026年1-2月", "CODEX单细胞数据预处理与baseline输入构建", "已完成"],
            ["2026年2月", "空间邻接建模与核心指标计算", "已完成半径20/30相关指标"],
            ["2026年3月", "原发灶与转移灶总体配对统计", "已完成，并识别整体异质性较大"],
            ["2026年3-4月", "R-only与L-only分层分析、中期报告整理", "已完成初步统计，正在补图表"],
            ["2026年5月", "结果可视化与论文初稿", "尚未完成"],
        ],
    )
    append_paragraph(
        cell,
        "结合pair_manifest.tsv可知，当前共建立16组原发灶-转移灶配对，其中R组7对、L组5对、QT组4对。考虑到课题主线是卵巢转移分析，现阶段正式主分析范围固定为P对L/R，QT仅作为补充分析对象，不纳入卵巢转移主结论。",
    )
    append_paragraph(
        cell,
        "这一调整也是对开题阶段“原发灶 vs 卵巢转移灶”研究对象定义的进一步收紧：中期分析不再泛化为“所有转移灶比较”，而是将开题中强调的卵巢转移问题落到L/R样本层面，使研究对象与题目表述保持一致。",
    )

    append_paragraph(cell, "2.已完成内容", bold=True, first_indent=False)
    append_paragraph(
        cell,
        "（1）已完成数据与研究对象定义。依据开题答辩阶段的数据说明文件与样本表，明确了本研究包含14例病人、31个标本，样本编码规则为P、L、R、QT，并在后续实验中统一了“P vs L/R”为主分析、“QT排除出卵巢转移主分析”的口径。",
        first_indent=False,
    )
    append_paragraph(
        cell,
        "（2）已完成baseline数据整理与样本级汇总。基于sample_summary.tsv与pair_manifest.tsv，建立了样本清单、病人配对关系以及肿瘤细胞、M1型巨噬细胞、M2型巨噬细胞的样本级统计信息。",
        first_indent=False,
    )
    append_paragraph(
        cell,
        "（3）已完成空间建模与指标计算。当前实验已围绕固定半径邻域构建空间指标体系，完成距离类、within-r类、局部互作强度类及巨噬细胞组成类指标的计算，并形成cohort_analysis.tsv和stratified_stats.tsv等结果文件。",
        first_indent=False,
    )
    append_paragraph(
        cell,
        "上述内容与开题报告4.3、4.4中的实施方案基本对应，说明开题阶段提出的“单细胞空间邻接图建模-空间指标计算-配对统计检验”主路线已经进入可复述、可展示、可继续扩展的执行状态。",
        first_indent=False,
    )
    append_paragraph(
        cell,
        "（4）已完成总体配对分析和分层统计。总体结果提示不同病人之间存在较大异质性，因此进一步对R-only和L-only进行了分层比较。现阶段最适合作为中期主结论展示的是R-only组中的Tumor-M1互作下降，以及M2相关指标的方向性变化。",
        first_indent=False,
    )

    r_enrich = get_metric(stratified_rows, "R-only", "enrich_30_Tumor_M1")
    r_within = get_metric(stratified_rows, "R-only", "within_30_M2")
    r_m2 = get_metric(stratified_rows, "R-only", "m2_frac_of_mac")
    l_enrich = get_metric(stratified_rows, "L-only", "enrich_30_Tumor_M1")
    l_within = get_metric(stratified_rows, "L-only", "within_30_M2")
    l_m2 = get_metric(stratified_rows, "L-only", "m2_frac_of_mac")

    append_paragraph(cell, "已完成的关键分层结果如下：", first_indent=False)
    add_simple_table(
        cell,
        ["分层", "指标", "配对数", "均值差", "中位差", "p值", "写法"],
        [
            ["R-only", "Tumor-M1富集(30)", format_num(r_enrich["n_pairs"]), format_num(r_enrich["diff_mean"]), format_num(r_enrich["diff_median"]), format_num(r_enrich["wilcoxon_p"]), "主结果：下降"],
            ["R-only", "M2邻域占比(30)", format_num(r_within["n_pairs"]), format_num(r_within["diff_mean"]), format_num(r_within["diff_median"]), format_num(r_within["wilcoxon_p"]), "辅助：重排"],
            ["R-only", "M2/巨噬比例", format_num(r_m2["n_pairs"]), format_num(r_m2["diff_mean"]), format_num(r_m2["diff_median"]), format_num(r_m2["wilcoxon_p"]), "辅助：上升"],
            ["L-only", "Tumor-M1富集(30)", format_num(l_enrich["n_pairs"]), format_num(l_enrich["diff_mean"]), format_num(l_enrich["diff_median"]), format_num(l_enrich["wilcoxon_p"]), "保守描述"],
            ["L-only", "M2邻域占比(30)", format_num(l_within["n_pairs"]), format_num(l_within["diff_mean"]), format_num(l_within["diff_median"]), format_num(l_within["wilcoxon_p"]), "保守描述"],
            ["L-only", "M2/巨噬比例", format_num(l_m2["n_pairs"]), format_num(l_m2["diff_mean"]), format_num(l_m2["diff_median"]), format_num(l_m2["wilcoxon_p"]), "保守描述"],
        ],
    )
    append_paragraph(
        cell,
        "从当前结果看，R-only组呈现出更清晰的方向一致性：Tumor-M1局部互作减弱，M2在巨噬细胞中的比例上升，同时M2在肿瘤周围的贴边比例下降。L-only组受样本量限制，暂不具备形成稳定主结论的条件。",
    )
    append_paragraph(
        cell,
        "从与开题预期成果的对应关系看，中期阶段已经拿到了“至少两类空间指标在原发灶与转移灶中的对比结果”，并且形成了可用于答辩陈述的主结果雏形；尚未完成的主要是开题中提到的“空间模式归纳”和更完整的图表化输出。",
    )

    append_paragraph(cell, "3.尚需完成内容", bold=True, first_indent=False)
    append_paragraph(
        cell,
        "（1）补充中期答辩所需图表。现阶段正文和统计表已基本具备，但仍需补充R-only代表病例的空间示意图、enrich_30_Tumor_M1主结果图、within_30_M2辅助图，以及与开题技术路线对应的流程图。",
        first_indent=False,
    )
    append_paragraph(
        cell,
        "（2）进一步完善结果解释与稳健性分析。当前半径20和30的结果已经形成初步趋势，但仍需检查参数选择与细胞筛选规则对结论稳定性的影响，使后续论文写作中的方法部分更加严谨。",
        first_indent=False,
    )
    append_paragraph(
        cell,
        "（3）完成中期答辩展示材料与论文初稿衔接。后续需要将现有实验结果整理为答辩PPT中的图表，并同步衔接到毕业论文初稿的方法与结果章节中。",
        first_indent=False,
    )

    append_paragraph(cell, "4.存在的问题", bold=True, first_indent=False)
    append_paragraph(
        cell,
        "（1）样本异质性较大。总体配对分析中不同病人之间差异明显，导致所有转移灶简单混合后难以得到稳定、易解释的总体结论。",
        first_indent=False,
    )
    append_paragraph(
        cell,
        "（2）不同转移部位不能混写。QT表示除卵巢外的其他转移灶，生物学含义不同于L/R，因此不能并入卵巢转移主分析；如果分组定义不严谨，会直接影响结论对象。",
        first_indent=False,
    )
    append_paragraph(
        cell,
        "（3）L-only组样本数相对较少，关键指标波动较大，现阶段尚不足以支撑与R-only的强比较。",
        first_indent=False,
    )
    append_paragraph(
        cell,
        "（4）当前图像级展示不足。虽然已经形成样本级和分层统计结果，但代表病例空间图和答辩用主结果图仍未补齐，影响中期答辩中的直观表达。",
        first_indent=False,
    )

    append_paragraph(cell, "5.解决措施", bold=True, first_indent=False)
    append_paragraph(
        cell,
        "（1）固定主分析口径。后续所有正式写作与图表均坚持“P vs L/R”为卵巢转移主分析，“QT仅作补充说明”，避免研究对象漂移。",
        first_indent=False,
    )
    append_paragraph(
        cell,
        "（2）以R-only结果作为中期阶段主线。现阶段围绕R-only中的enrich_30_Tumor_M1、within_30_M2和m2_frac_of_mac组织图表与叙述，L-only只保守描述，QT不写主结论。",
        first_indent=False,
    )
    append_paragraph(
        cell,
        "（3）补齐可视化与说明材料。优先补做1-2个R-only代表病例的空间图、关键指标配对图和技术路线图，使中期报告与答辩展示相互支撑。",
        first_indent=False,
    )
    append_paragraph(
        cell,
        "（4）继续完善参数与规则说明。对半径参数、细胞筛选规则、指标定义和统计口径进行统一整理，在后续论文写作中形成更规范的方法描述，必要时补充稳健性检验作为附加分析。",
        first_indent=False,
    )

    # Remove the first empty paragraph created by clear_cell if still empty.
    if cell.paragraphs and not cell.paragraphs[0].text.strip():
        p = cell.paragraphs[0]._element
        p.getparent().remove(p)


def main():
    doc = Document(str(TEMPLATE_PATH))
    if doc.paragraphs:
        doc.paragraphs[0].text = "中国地质大学（北京）本科毕业设计（论文）中期报告"
        if doc.paragraphs[0].runs:
            for run in doc.paragraphs[0].runs:
                set_run_font(run, size=15, bold=True)
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.tables[0]
    fill_meta(table)
    build_body(table.cell(4, 0))

    doc.save(str(REPORT_PATH))


if __name__ == "__main__":
    main()
